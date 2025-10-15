import sys
import cv2
import pytesseract
from docx import Document
from docx.shared import Inches
from PIL import Image

input_path = sys.argv[1]
output_path = sys.argv[2]

# Load image
image = cv2.imread(input_path)

# Create new Word document
doc = Document()

# Add the full image as reference
doc.add_picture(input_path, width=Inches(6))

# Extract layout-aware text (with coordinates)
data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)

current_y = 0
for i in range(len(data["text"])):
    if data["text"][i].strip() != "":
        text = data["text"][i]
        y = data["top"][i]

        # Group text lines by Y-position
        if abs(y - current_y) > 25:
            doc.add_paragraph("")
        doc.add_run(text + " ")
        current_y = y

# Save the document
doc.save(output_path)
