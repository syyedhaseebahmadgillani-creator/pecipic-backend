#!/usr/bin/env python3
# python/layout_ocr.py
# Usage: python3 layout_ocr.py input_image_path output_docx_path

import sys
import os
import cv2
import numpy as np
import pytesseract
from docx import Document
from docx.shared import Inches
from PIL import Image

def detect_blocks(image_path):
    img = cv2.imread(image_path)
    if img is None:
        raise RuntimeError("Failed to open image: " + image_path)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    thr = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C,
                                cv2.THRESH_BINARY_INV, 15, 10)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (30, 15))
    dilated = cv2.dilate(thr, kernel, iterations=2)
    contours, _ = cv2.findContours(dilated, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    blocks = []
    h, w = gray.shape
    for cnt in contours:
        x, y, bw, bh = cv2.boundingRect(cnt)
        if bw < 50 or bh < 20:
            continue
        pad = 5
        x1 = max(0, x - pad)
        y1 = max(0, y - pad)
        x2 = min(w, x + bw + pad)
        y2 = min(h, y + bh + pad)
        blocks.append((x1, y1, x2, y2))
    blocks = sorted(blocks, key=lambda b: (b[1], b[0]))
    return blocks

def ocr_block(image_path, bbox):
    x1, y1, x2, y2 = bbox
    img = Image.open(image_path).convert("RGB")
    crop = img.crop((x1, y1, x2, y2))
    text = pytesseract.image_to_string(crop, lang='eng')
    return text.strip()

def build_docx(image_path, blocks, output_path):
    doc = Document()
    doc.add_picture(image_path, width=Inches(6))
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("——— Extracted Blocks ———")
    run.bold = True
    for bbox in blocks:
        text = ocr_block(image_path, bbox)
        if text:
            p = doc.add_paragraph()
            p.add_run(text)
            doc.add_paragraph("")
        else:
            doc.add_paragraph("")
    doc.add_paragraph("\nConverted by Pecipic.ai").italic = True
    doc.save(output_path)

def main():
    if len(sys.argv) < 3:
        print("Usage: layout_ocr.py input_image output_docx", file=sys.stderr)
        sys.exit(2)
    input_image = sys.argv[1]
    output_docx = sys.argv[2]
    blocks = detect_blocks(input_image)
    build_docx(input_image, blocks, output_docx)
    print("OK")

if __name__ == "__main__":
    main()
